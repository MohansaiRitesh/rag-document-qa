"""
Document Processing Module

This module handles:
1. Loading documents (PDF, DOCX, TXT)
2. Splitting documents into chunks
3. Extracting metadata

"""

from typing import List, Dict, Tuple, Union
from pathlib import Path
import re
import numpy as np
import time
import uuid
import os
import io
import base64
import PyPDF2
import fitz  # PyMuPDF
from docx import Document
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
from groq import Groq


class DocumentProcessor:
    """
    Handles all document processing operations
    """
    
    def __init__(
        self, 
        chunk_size: int = 1000, 
        chunk_overlap: int = 200,
        chunking_strategy: str = "recursive",
        semantic_threshold_alpha: float = 1.0,
        semantic_max_chunk_size: int = 1500,
        embedding_model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
        parent_chunk_size: int = 1500,
        parent_chunk_overlap: int = 200,
        child_chunk_size: int = 300,
        child_chunk_overlap: int = 50,
        api_key: str = None
    ):
        """
        Initialize the processor with chunking parameters
        
        Args:
            chunk_size: Maximum characters per chunk
            chunk_overlap: Characters to overlap between chunks
            chunking_strategy: "recursive", "semantic" or "hierarchical"
            semantic_threshold_alpha: Scale factor for standard deviation thresholding
            semantic_max_chunk_size: Fallback safety limit for semantic chunks
            embedding_model_name: Name of model for embedding sentences in semantic chunking
            parent_chunk_size: Characters per parent chunk (hierarchical)
            parent_chunk_overlap: Characters overlap between parent chunks
            child_chunk_size: Characters per child chunk (hierarchical)
            child_chunk_overlap: Characters overlap between child chunks
            api_key: Groq API key for vision captioning
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.chunking_strategy = chunking_strategy
        self.semantic_threshold_alpha = semantic_threshold_alpha
        self.semantic_max_chunk_size = semantic_max_chunk_size
        self.embedding_model_name = embedding_model_name
        self.embedding_model = None
        self.parent_chunk_size = parent_chunk_size
        self.parent_chunk_overlap = parent_chunk_overlap
        self.child_chunk_size = child_chunk_size
        self.child_chunk_overlap = child_chunk_overlap
        
        self.groq_client = Groq(api_key=api_key) if api_key else None
        
        # Recursive splitter (standard)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]  # Priority order
        )
        
        # Parent splitter (hierarchical)
        self.parent_splitter = RecursiveCharacterTextSplitter(
            chunk_size=parent_chunk_size,
            chunk_overlap=parent_chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Child splitter (hierarchical)
        self.child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=child_chunk_size,
            chunk_overlap=child_chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

    def _get_embedding_model(self) -> SentenceTransformer:
        """Lazily load sentence-transformer model"""
        if self.embedding_model is None:
            print(f"[INFO] Loading embedding model for Semantic Chunking: {self.embedding_model_name}")
            self.embedding_model = SentenceTransformer(self.embedding_model_name)
            print("[INFO] Embedding model loaded successfully for Semantic Chunking")
        return self.embedding_model

    def _split_into_sentences(self, text: str) -> List[str]:
        """
        Split text into sentences or paragraphs/headers.
        """
        # First split by newline paragraphs to capture headers/sections properly
        paragraphs = text.split('\n')
        
        sentences = []
        sentence_endings = re.compile(r'(?<=[\.\?\!])\s+')
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            # Split paragraph into sentences
            splits = sentence_endings.split(para)
            for s in splits:
                s_strip = s.strip()
                if s_strip:
                    sentences.append(s_strip)
        return sentences
    
    def load_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text as string
        """
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    # Add page number for citation
                    text += f"\n[Page {page_num + 1}]\n{page_text}"
            return text
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")

    def caption_image(self, image_bytes: bytes) -> str:
        """Call Groq to caption the image"""
        if not self.groq_client:
            raise ValueError("Groq client not initialized in DocumentProcessor")
            
        base64_image = base64.b64encode(image_bytes).decode('utf-8')
        
        response = self.groq_client.chat.completions.create(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "Identify if this image is a chart, graph, table, or general illustration. "
                                "Describe it in detail for a RAG search engine: transcribe any visible text, "
                                "describe data trends, labels, or relationships, and summarize the key information shown. "
                                "Keep the description factual and comprehensive."
                            )
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}"
                            }
                        }
                    ]
                }
            ],
            temperature=0.2,
            max_tokens=512
        )
        return response.choices[0].message.content.strip()

    def extract_and_caption_pdf_images(self, file_path: str, metadata: Dict) -> List[LangchainDocument]:
        """Extract all embedded images in a PDF and create descriptive chunk documents for them"""
        extracted_docs = []
        try:
            doc = fitz.open(file_path)
        except Exception as e:
            print(f"[WARNING] Could not open PDF with PyMuPDF for image extraction: {str(e)}")
            return []

        from config import get_settings
        settings = get_settings()
        extracted_images_dir = settings.extracted_images_dir

        for page_num in range(len(doc)):
            page = doc[page_num]
            try:
                image_list = page.get_images(full=True)
            except Exception as e:
                print(f"[WARNING] Could not get page {page_num+1} images: {str(e)}")
                continue

            for img_idx, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]

                    # Skip tiny decorative/spacer icons under 2KB
                    if len(image_bytes) < 2048:
                        continue

                    # Generate unique name
                    image_uuid = str(uuid.uuid4())
                    image_filename = f"{image_uuid}.{image_ext}"
                    save_path = os.path.join(extracted_images_dir, image_filename)

                    with open(save_path, "wb") as f:
                        f.write(image_bytes)

                    # Call Groq to caption
                    if self.groq_client:
                        try:
                            caption = self.caption_image(image_bytes)
                            print(f"[INFO] Generated caption for visual element: '{caption[:80]}...'")
                        except Exception as caption_err:
                            print(f"[WARNING] Image captioning failed: {str(caption_err)}")
                            caption = "Image extracted but VLM captioning failed."
                    else:
                        caption = "Image extracted but Groq API key is not configured."

                    doc_obj = LangchainDocument(
                        page_content=f"[Visual Content Description]: {caption}",
                        metadata={
                            **metadata,
                            "chunk_id": f"img_{page_num+1}_{img_idx}",
                            "is_image": True,
                            "image_path": image_filename,
                            "page_number": page_num + 1
                        }
                    )
                    extracted_docs.append(doc_obj)
                except Exception as img_err:
                    print(f"[WARNING] Skipping image extraction on page {page_num+1} index {img_idx}: {str(img_err)}")
                    continue

        return extracted_docs
    
    def load_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text as string
        """
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def load_txt(self, file_path: str) -> str:
        """
        Read text file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            File content as string
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")
    
    def load_document(self, file_path: str) -> str:
        """
        Load document based on file extension
        
        Args:
            file_path: Path to document
            
        Returns:
            Extracted text
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        loaders = {
            '.pdf': self.load_pdf,
            '.docx': self.load_docx,
            '.txt': self.load_txt
        }
        
        if extension not in loaders:
            raise ValueError(f"Unsupported file type: {extension}")
        
        return loaders[extension](file_path)
    
    def create_semantic_chunks(
        self, 
        text: str, 
        metadata: Dict = None,
        semantic_threshold_alpha: float = None
    ) -> List[LangchainDocument]:
        """
        Split text using semantic similarity between adjacent sentences.
        """
        if metadata is None:
            metadata = {}
            
        sentences = self._split_into_sentences(text)
        if not sentences:
            return []
            
        # If there is only one sentence, return it as a single chunk
        if len(sentences) == 1:
            return [LangchainDocument(page_content=sentences[0], metadata={**metadata, "chunk_id": 0, "total_chunks": 1})]
            
        # Get embeddings
        model = self._get_embedding_model()
        embeddings = model.encode(sentences, show_progress_bar=False, convert_to_numpy=True)
        
        # Calculate cosine distances between adjacent sentences
        distances = []
        for i in range(len(embeddings) - 1):
            vec1 = embeddings[i]
            vec2 = embeddings[i + 1]
            
            norm1 = np.linalg.norm(vec1)
            norm2 = np.linalg.norm(vec2)
            
            if norm1 > 0 and norm2 > 0:
                cosine_sim = np.dot(vec1, vec2) / (norm1 * norm2)
            else:
                cosine_sim = 0.0
                
            distances.append(1.0 - cosine_sim)
            
        # Determine threshold: mean + alpha * std
        alpha = semantic_threshold_alpha if semantic_threshold_alpha is not None else self.semantic_threshold_alpha
        if distances:
            mean_dist = np.mean(distances)
            std_dist = np.std(distances)
            threshold = mean_dist + alpha * std_dist
        else:
            threshold = 0.5
            
        # Group sentences into chunks
        chunks = []
        current_chunk_sentences = []
        current_chunk_char_count = 0
        
        for i, sentence in enumerate(sentences):
            current_chunk_sentences.append(sentence)
            current_chunk_char_count += len(sentence)
            
            # Decide if we split after sentence i
            should_split = False
            
            # Check threshold (only if there is a next sentence)
            if i < len(distances):
                if distances[i] > threshold:
                    should_split = True
            
            # Check maximum chunk size limit to prevent too large chunks
            if current_chunk_char_count >= self.semantic_max_chunk_size:
                should_split = True
                
            if should_split and current_chunk_sentences:
                chunk_text = " ".join(current_chunk_sentences)
                chunks.append(chunk_text)
                current_chunk_sentences = []
                current_chunk_char_count = 0
                
        # Add any remaining sentences
        if current_chunk_sentences:
            chunk_text = " ".join(current_chunk_sentences)
            chunks.append(chunk_text)
            
        # Map to LangchainDocument structures
        documents = []
        for i, chunk in enumerate(chunks):
            doc = LangchainDocument(
                page_content=chunk,
                metadata={
                    **metadata,
                    "chunk_id": i,
                    "total_chunks": len(chunks)
                }
            )
            documents.append(doc)
            
        return documents

    def create_hierarchical_chunks(
        self, 
        text: str, 
        metadata: Dict = None
    ) -> Tuple[List[LangchainDocument], Dict[str, Dict]]:
        """
        Split document into parent and child chunks.
        
        Returns:
            Tuple of:
            - List of LangchainDocument objects for child chunks (to embed and search)
            - Dict of parent_id -> {"text": parent_text, "metadata": parent_metadata} (to save in lookup)
        """
        if metadata is None:
            metadata = {}
            
        parent_chunks = self.parent_splitter.split_text(text)
        child_documents = []
        parent_mappings = {}
        
        child_count = 0
        for i, parent_text in enumerate(parent_chunks):
            parent_id = str(uuid.uuid4())
            
            # Save parent metadata
            parent_metadata = {
                **metadata,
                "parent_id": parent_id,
                "chunk_id": i,
                "total_parent_chunks": len(parent_chunks),
                "is_parent": True
            }
            parent_mappings[parent_id] = {
                "text": parent_text,
                "metadata": parent_metadata
            }
            
            # Split parent text into child chunks
            child_chunks = self.child_splitter.split_text(parent_text)
            for j, child_text in enumerate(child_chunks):
                child_metadata = {
                    **metadata,
                    "parent_id": parent_id,
                    "chunk_id": child_count,  # Added for consistency across strategies
                    "child_id": child_count,
                    "local_child_id": j,
                    "total_local_child_chunks": len(child_chunks),
                    "is_child": True
                }
                
                doc = LangchainDocument(
                    page_content=child_text,
                    metadata=child_metadata
                )
                child_documents.append(doc)
                child_count += 1
                
        return child_documents, parent_mappings

    def create_chunks(
        self, 
        text: str, 
        metadata: Dict = None,
        chunking_strategy: str = None,
        semantic_threshold_alpha: float = None
    ) -> Union[List[LangchainDocument], Tuple[List[LangchainDocument], Dict[str, Dict]]]:
        """
        Split text into chunks with metadata
        
        Args:
            text: Document text to split
            metadata: Document metadata (source, title, etc.)
            chunking_strategy: Optional strategy override ("recursive", "semantic" or "hierarchical")
            semantic_threshold_alpha: Optional threshold scale override
            
        Returns:
            - List of LangchainDocument chunks for recursive/semantic strategies
            - Tuple of (child_docs, parent_mappings) for hierarchical strategy
        """
        strategy = chunking_strategy or self.chunking_strategy
        if strategy == "semantic":
            return self.create_semantic_chunks(text, metadata, semantic_threshold_alpha)
        elif strategy == "hierarchical":
            return self.create_hierarchical_chunks(text, metadata)
        else:
            if metadata is None:
                metadata = {}
            
            chunks = self.text_splitter.split_text(text)
            
            documents = []
            for i, chunk in enumerate(chunks):
                doc = LangchainDocument(
                    page_content=chunk,
                    metadata={
                        **metadata,
                        "chunk_id": i,
                        "total_chunks": len(chunks)
                    }
                )
                documents.append(doc)
            return documents
    
    def process_document(
        self, 
        file_path: str,
        chunking_strategy: str = None,
        semantic_threshold_alpha: float = None
    ) -> Union[List[LangchainDocument], Tuple[List[LangchainDocument], Dict[str, Dict]]]:
        """
        Complete pipeline: Load → Chunk → Add Metadata (including visual content extraction)
        
        Args:
            file_path: Path to document
            chunking_strategy: Optional strategy override ("recursive", "semantic" or "hierarchical")
            semantic_threshold_alpha: Optional threshold scale override
            
        Returns:
            List of child chunks or Tuple of (child_chunks, parent_mappings)
        """
        # Load document
        text = self.load_document(file_path)
        
        # Extract filename for metadata
        filename = Path(file_path).name
        path = Path(file_path)
        
        # Create metadata
        metadata = {
            "source": filename,
            "file_path": file_path,
            "file_type": path.suffix.lower().lstrip("."),
            "file_size": path.stat().st_size if path.exists() else 0,
            "uploaded_at": time.time()
        }
        
        # Create text chunks
        result = self.create_chunks(
            text, 
            metadata, 
            chunking_strategy=chunking_strategy,
            semantic_threshold_alpha=semantic_threshold_alpha
        )
        
        # If it is a PDF, extract and caption images
        if path.suffix.lower() == ".pdf":
            print(f"[INFO] Scanning {filename} for embedded visual elements...")
            visual_docs = self.extract_and_caption_pdf_images(file_path, metadata)
            if visual_docs:
                print(f"[SUCCESS] Extracted and captioned {len(visual_docs)} visual elements from {filename}")
                if isinstance(result, tuple):
                    # Hierarchical mode: append to child chunks list
                    result[0].extend(visual_docs)
                else:
                    # Recursive/Semantic mode: extend chunks list
                    result.extend(visual_docs)
            else:
                print(f"[INFO] No visual elements extracted from {filename}")
        
        # Log count of child chunks created
        child_count = len(result[0]) if isinstance(result, tuple) else len(result)
        print(f"[SUCCESS] Processed {filename}: {child_count} total chunks (including visual chunks) created")
        return result


# Example usage (for testing)
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Test with a sample text
    sample_text = "This is a sample document. " * 100
    chunks = processor.create_chunks(sample_text, {"source": "test.txt"})
    
    print(f"Created {len(chunks)} chunks")
    print(f"First chunk preview: {chunks[0].page_content[:100]}...")
